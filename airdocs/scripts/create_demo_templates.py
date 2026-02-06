# AirDocs - Demo Template Generator
# =========================================
#
# This script creates demo template files for the application.
# Run it once after installation to set up initial templates.

import sys
from pathlib import Path

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))


def create_word_invoice_template(output_path: Path):
    """Create a demo Word invoice template using docxtpl."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Header
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run("СЧЕТ № {{ invoice_number }}")
    header_run.bold = True
    header_run.font.size = Pt(16)

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run("от {{ invoice_date }}")

    doc.add_paragraph()

    # Seller info
    doc.add_paragraph("Поставщик: {{ seller_name }}")
    doc.add_paragraph("ИНН/КПП: {{ seller_inn }} / {{ seller_kpp }}")
    doc.add_paragraph("Адрес: {{ seller_address }}")
    doc.add_paragraph()

    # Buyer info
    doc.add_paragraph("Покупатель: {{ buyer_name }}")
    doc.add_paragraph("ИНН/КПП: {{ buyer_inn }} / {{ buyer_kpp }}")
    doc.add_paragraph("Адрес: {{ buyer_address }}")
    doc.add_paragraph()

    # AWB reference
    doc.add_paragraph("AWB: {{ awb_number }}")
    doc.add_paragraph("Маршрут: {{ departure_airport }} - {{ destination_airport }}")
    doc.add_paragraph()

    # Table
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    # Header row
    headers = ["№", "Наименование услуги", "Кол-во", "Цена", "Сумма"]
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True

    # Data row (template)
    row = table.add_row()
    row.cells[0].text = "1"
    row.cells[1].text = "{{ service_description }}"
    row.cells[2].text = "{{ quantity }}"
    row.cells[3].text = "{{ unit_price }}"
    row.cells[4].text = "{{ total_amount }}"

    doc.add_paragraph()

    # Totals
    totals_para = doc.add_paragraph()
    totals_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    totals_para.add_run("Итого: {{ total_amount }} руб.").bold = True

    doc.add_paragraph()

    total_words = doc.add_paragraph()
    total_words.add_run("Сумма прописью: {{ total_amount_words }}")

    doc.add_paragraph()
    doc.add_paragraph()

    # Signatures
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.rows[0].cells[0].text = "Поставщик: _____________ / {{ seller_signatory }} /"
    sig_table.rows[0].cells[1].text = "Покупатель: _____________ / {{ buyer_signatory }} /"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Created: {output_path}")


def create_word_upd_template(output_path: Path):
    """Create a demo UPD (Universal Transfer Document) template."""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("УНИВЕРСАЛЬНЫЙ ПЕРЕДАТОЧНЫЙ ДОКУМЕНТ")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Document info
    doc.add_paragraph("Номер: {{ upd_number }}")
    doc.add_paragraph("Дата: {{ upd_date }}")
    doc.add_paragraph()

    # Seller
    doc.add_paragraph("Продавец: {{ seller_name }}")
    doc.add_paragraph("ИНН/КПП: {{ seller_inn }} / {{ seller_kpp }}")
    doc.add_paragraph("Адрес: {{ seller_address }}")
    doc.add_paragraph()

    # Buyer
    doc.add_paragraph("Покупатель: {{ buyer_name }}")
    doc.add_paragraph("ИНН/КПП: {{ buyer_inn }} / {{ buyer_kpp }}")
    doc.add_paragraph("Адрес: {{ buyer_address }}")
    doc.add_paragraph()

    # Service details
    doc.add_paragraph("Основание: Договор № {{ contract_number }} от {{ contract_date }}")
    doc.add_paragraph("AWB: {{ awb_number }}")
    doc.add_paragraph()

    # Table
    table = doc.add_table(rows=2, cols=6)
    table.style = "Table Grid"

    headers = ["№", "Наименование", "Ед.", "Кол-во", "Цена", "Сумма"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "{{ service_description }}"
    table.rows[1].cells[2].text = "шт"
    table.rows[1].cells[3].text = "{{ quantity }}"
    table.rows[1].cells[4].text = "{{ unit_price }}"
    table.rows[1].cells[5].text = "{{ total_amount }}"

    doc.add_paragraph()
    doc.add_paragraph("Итого: {{ total_amount }} руб., в т.ч. НДС {{ vat_amount }} руб.")
    doc.add_paragraph()

    # Signatures
    doc.add_paragraph("Товар (услуга) передал: _____________ / {{ seller_signatory }} /")
    doc.add_paragraph("Товар (услугу) принял: _____________ / {{ buyer_signatory }} /")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Created: {output_path}")


def create_word_act_template(output_path: Path):
    """Create a demo Act of services template."""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("АКТ № {{ act_number }}")
    run.bold = True
    run.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("выполненных работ (оказанных услуг)")

    doc.add_paragraph()
    doc.add_paragraph("г. Москва                                                           {{ act_date }}")
    doc.add_paragraph()

    doc.add_paragraph("Исполнитель: {{ seller_name }}, ИНН {{ seller_inn }}")
    doc.add_paragraph("Заказчик: {{ buyer_name }}, ИНН {{ buyer_inn }}")
    doc.add_paragraph()
    doc.add_paragraph("Основание: {{ contract_basis }}")
    doc.add_paragraph("AWB: {{ awb_number }}")
    doc.add_paragraph()

    # Services table
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"

    headers = ["№", "Наименование услуги", "Сумма без НДС", "Сумма с НДС"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "{{ service_description }}"
    table.rows[1].cells[2].text = "{{ amount_without_vat }}"
    table.rows[1].cells[3].text = "{{ total_amount }}"

    doc.add_paragraph()
    doc.add_paragraph("Итого оказано услуг на сумму: {{ total_amount }} руб.")
    doc.add_paragraph("В том числе НДС: {{ vat_amount }} руб.")
    doc.add_paragraph()

    doc.add_paragraph("Вышеперечисленные услуги выполнены полностью и в срок.")
    doc.add_paragraph("Заказчик претензий по объему, качеству и срокам оказания услуг не имеет.")
    doc.add_paragraph()

    doc.add_paragraph("Исполнитель: _____________ / {{ seller_signatory }} /")
    doc.add_paragraph()
    doc.add_paragraph("Заказчик: _____________ / {{ buyer_signatory }} /")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Created: {output_path}")


def create_excel_registry_template(output_path: Path):
    """Create a demo Excel registry template for 1C."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

    wb = Workbook()
    ws = wb.active
    ws.title = "Реестр"

    # Styles
    header_font = Font(bold=True, size=11)
    header_fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )
    center_align = Alignment(horizontal="center", vertical="center")

    # Title
    ws.merge_cells("A1:J1")
    ws["A1"] = "РЕЕСТР ДОКУМЕНТОВ"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A1"].alignment = center_align

    ws.merge_cells("A2:J2")
    ws["A2"] = "Дата формирования: {{ registry_date }}"
    ws["A2"].alignment = center_align

    # Headers (row 4)
    headers = [
        ("A", "№ п/п"),
        ("B", "AWB"),
        ("C", "Дата"),
        ("D", "Отправитель"),
        ("E", "Получатель"),
        ("F", "Маршрут"),
        ("G", "Мест"),
        ("H", "Вес, кг"),
        ("I", "Сумма, руб"),
        ("J", "Статус"),
    ]

    for col, header in headers:
        cell = ws[f"{col}4"]
        cell.value = header
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_align

    # Template row (row 5) - will be repeated for each record
    ws["A5"] = "{{ loop.index }}"
    ws["B5"] = "{{ item.awb_number }}"
    ws["C5"] = "{{ item.date }}"
    ws["D5"] = "{{ item.shipper_name }}"
    ws["E5"] = "{{ item.consignee_name }}"
    ws["F5"] = "{{ item.route }}"
    ws["G5"] = "{{ item.pieces }}"
    ws["H5"] = "{{ item.weight_kg }}"
    ws["I5"] = "{{ item.total_amount }}"
    ws["J5"] = "{{ item.status }}"

    for col in "ABCDEFGHIJ":
        ws[f"{col}5"].border = thin_border

    # Column widths
    widths = {"A": 8, "B": 15, "C": 12, "D": 25, "E": 25, "F": 15, "G": 8, "H": 10, "I": 12, "J": 12}
    for col, width in widths.items():
        ws.column_dimensions[col].width = width

    # Summary row
    ws["A7"] = "ИТОГО:"
    ws["A7"].font = header_font
    ws["I7"] = "{{ total_sum }}"
    ws["I7"].font = header_font

    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(output_path))
    print(f"Created: {output_path}")


def create_pdf_awb_blank(output_path: Path):
    """Create a simple AWB blank PDF placeholder."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas
    from reportlab.lib.colors import lightgrey, black

    output_path.parent.mkdir(parents=True, exist_ok=True)

    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4

    # Title
    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(width / 2, height - 30 * mm, "AIR WAYBILL")

    # AWB number box
    c.setStrokeColor(black)
    c.setFillColor(lightgrey)
    c.rect(width - 100 * mm, height - 50 * mm, 80 * mm, 15 * mm, fill=1)
    c.setFillColor(black)
    c.setFont("Helvetica", 8)
    c.drawString(width - 95 * mm, height - 42 * mm, "AWB Number")

    # Shipper box
    c.setFillColor(lightgrey)
    c.rect(20 * mm, height - 100 * mm, 80 * mm, 40 * mm, fill=1)
    c.setFillColor(black)
    c.setFont("Helvetica", 8)
    c.drawString(22 * mm, height - 65 * mm, "Shipper's Name and Address")

    # Consignee box
    c.setFillColor(lightgrey)
    c.rect(20 * mm, height - 150 * mm, 80 * mm, 40 * mm, fill=1)
    c.setFillColor(black)
    c.drawString(22 * mm, height - 115 * mm, "Consignee's Name and Address")

    # Airport boxes
    c.setFillColor(lightgrey)
    c.rect(20 * mm, height - 180 * mm, 40 * mm, 20 * mm, fill=1)
    c.rect(70 * mm, height - 180 * mm, 40 * mm, 20 * mm, fill=1)
    c.setFillColor(black)
    c.drawString(22 * mm, height - 165 * mm, "Airport of Departure")
    c.drawString(72 * mm, height - 165 * mm, "Airport of Destination")

    # Flight info
    c.setFillColor(lightgrey)
    c.rect(120 * mm, height - 180 * mm, 30 * mm, 20 * mm, fill=1)
    c.rect(160 * mm, height - 180 * mm, 30 * mm, 20 * mm, fill=1)
    c.setFillColor(black)
    c.drawString(122 * mm, height - 165 * mm, "Flight/Date")
    c.drawString(162 * mm, height - 165 * mm, "Flight/Date")

    # Cargo details table
    c.setFillColor(lightgrey)
    c.rect(20 * mm, height - 250 * mm, 170 * mm, 60 * mm, fill=1)
    c.setFillColor(black)
    c.drawString(22 * mm, height - 195 * mm, "No. of Pieces")
    c.drawString(60 * mm, height - 195 * mm, "Gross Weight")
    c.drawString(100 * mm, height - 195 * mm, "Rate Class")
    c.drawString(140 * mm, height - 195 * mm, "Total")

    # Description box
    c.setFillColor(lightgrey)
    c.rect(20 * mm, height - 320 * mm, 170 * mm, 60 * mm, fill=1)
    c.setFillColor(black)
    c.drawString(22 * mm, height - 265 * mm, "Nature and Quantity of Goods")

    # Footer note
    c.setFont("Helvetica", 6)
    c.drawString(20 * mm, 20 * mm, "This is a DEMO AWB blank for testing purposes only")

    c.save()
    print(f"Created: {output_path}")


def main():
    """Create all demo templates."""
    print("Creating demo templates...")
    print("=" * 50)

    base_path = Path(__file__).parent.parent / "templates"

    # Word templates
    create_word_invoice_template(base_path / "word" / "invoice.docx")
    create_word_upd_template(base_path / "word" / "upd.docx")
    create_word_act_template(base_path / "word" / "act.docx")

    # Excel templates
    create_excel_registry_template(base_path / "excel" / "registry_1c.xlsx")

    # PDF templates
    create_pdf_awb_blank(base_path / "pdf" / "awb_blank.pdf")

    print("=" * 50)
    print("All demo templates created successfully!")


if __name__ == "__main__":
    main()
